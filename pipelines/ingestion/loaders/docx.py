# pipelines/ingestion/loaders/docx.py
import docx
import io

def parse_docx_bytes(file_bytes: bytes, filename: str):
    """
    Parses a .docx file from memory bytes.
    Extracts text, headers, and simple tables.
    """
    file_stream = io.BytesIO(file_bytes)
    doc = docx.Document(file_stream)
    
    full_text = []
    
    # 1. Extract Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # 2. Extract Tables (Simple text conversion)
    # For complex tables, we would convert to HTML or Markdown
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    text_content = "\n\n".join(full_text)
    
    metadata = {
        "filename": filename,
        "type": "docx",
        "char_count": len(text_content)
    }
    
    return text_content, metadata